### ----------------------------
### TFBS Interdistances program
### ----------------------------

'''
This program allows to localize transcription factor binding sites on sequences such as promoters.
You need a matrix with frequency values and fasta sequences (bound sequences (i.e. peaks)).
This program was written by Adrien Bessy, Arnaud Stigliani and Francois Parcy, and was inspired by Morpheus program written by Eugenio Gomez Minguet
and (4-scores.py and CalculateScoreAboveTh_and_Interdistances.py) programs written by Laura Gregoire.
'''
import os,sys,inspect
sys.path.append('../get_interdistances/')
currentdir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
parentdir = os.path.dirname(currentdir)
sys.path.insert(0,parentdir) 
from interdistances_functions import *
import numpy as np
from Bio import SeqIO
import time
import sys 
from operator import truediv
import operator
from collections import Counter
import matplotlib.pyplot as plt
from matplotlib.offsetbox import AnchoredOffsetbox, TextArea, HPacker, VPacker
import matplotlib.patches as mpatches
from matplotlib import pylab
import types
import argparse
import logging
from optparse import OptionParser
from scipy import stats
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt


parser = argparse.ArgumentParser()                                               

parser.add_argument("--factor", "-fac", type=str, default= "ARF2")
parser.add_argument("--pseudoCount", "-pc",type=float,default= 0.001)
parser.add_argument("--threshold", "-th",nargs='+',type=int,default= -12)
parser.add_argument("--Interdistance_maxValue", "-maxInter", type=int,default= 20)
args = parser.parse_args()

#python get_bs_localizationsV2.py -fac "ARF5_OMalley_extended3Prime" -pc 0.001 -maxInter 20 -th -12
#python get_bs_localizationsV2.py -fac "ARF2_MEME1500" -pc 0.001 -maxInter 20 -th -12

factorTranscription = args.factor
pseudoCount = args.pseudoCount
threshold = args.threshold
Interdistance_maxValue = args.Interdistance_maxValue

                    
###################Parameters we can change#################

if factorTranscription == "ARF2_21bases" :
	FastaFile = "DR8_tested_in_EMSA.fas" 
	MatrixFile = "../matrix/ARF2_MEME_Matrix1_600FirstSeq.txt" 
	matrixType = "freq" 

if factorTranscription == "ARF2_MEME1500" :
	FastaFile = "fas_files/test12.fas" 
	MatrixFile = "../matrices/ARF2_matrix1_MEME_1500FirstSeq.txt" 
	matrixType = "freq" 	
	
if factorTranscription == "ARF2" :
	FastaFile = "fas_files/DR2.fas" 
	MatrixFile = "../matrices/ARF2_OMalley_matrixC.txt" 
	matrixType = "freq" 
	
if factorTranscription == "ARF5_OMalley_extended3Prime" :
	FastaFile = "fas_files/test12.fas" 
	MatrixFile = "../matrices/ARF5_allSeq_3prime_freq_pasteTo'OMalley.txt" 
	matrixType = "freq"	
	
if factorTranscription == "ARF5" :
	FastaFile = "fas_files/DR2.fas" 
	MatrixFile = "../matrices/ARF5_OMalley_matrixC.txt" 
	matrixType = "freq" 

############################################################################### 
	
document = Document()
document.add_heading("Localization of the "+factorTranscription+" binding sites on "+FastaFile, 0)	
	
def interdistance_calcul(DR,ER,IR,good_score_positions) :
	for first in range(0,len(good_score_positions)-1) :
		firstSubSeq = good_score_positions[first]
			
		for second in range(first+1,len(good_score_positions)) :
			secondSubSeq = good_score_positions[second]

			'''
			Here we do different substractions according to we get a Direct Repeat (DR), an Inverted Repeat (IR) or an Everted Repeat (ER).
			Because In the litterature, The interdistance calculation was made between subsequences from our matrix
			and not between the whole sequences from our matrix.
			So according to your specific interdistance calculations you can change these lines.
			'''
			
			if factorTranscription == "ARF2_21bases" :
				if firstSubSeq[1] == ">" and secondSubSeq[1] == ">" :
					d = ( int(secondSubSeq[0]) + 1 ) - ( int(firstSubSeq[0]) + 7)
					if Interdistance_maxValue >= d >= 0 :
						DR.append(["DR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == "<" and secondSubSeq[1] == "<" :
					d = ( int(secondSubSeq[0]) + 1 ) - ( int(firstSubSeq[0]) + 7)
					if Interdistance_maxValue >= d >= 0 :
						DR.append(["DR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == ">" and secondSubSeq[1] == "<" :
					d = ( int(secondSubSeq[0]) + 7 ) - ( int(firstSubSeq[0]) + 7 )
					if Interdistance_maxValue >= d >= 0 :
						ER.append(["ER "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == "<" and secondSubSeq[1] == ">" :
					d = ( int(secondSubSeq[0]) + 20 ) - ( int(firstSubSeq[0]) + 1 )
					if Interdistance_maxValue >= d >= 0 :
						IR.append(["IR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])

			
			if factorTranscription == "ARF2_MEME1500" :
				if firstSubSeq[1] == ">" and secondSubSeq[1] == ">" :
					d = ( int(secondSubSeq[0])  ) - ( int(firstSubSeq[0]) + 6 )
					if Interdistance_maxValue >= d >= 0 :
						DR.append(["DR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == "<" and secondSubSeq[1] == "<" :
					d = ( int(secondSubSeq[0]) ) - ( int(firstSubSeq[0]) + 6 )
					if Interdistance_maxValue >= d >= 0 :
						DR.append(["DR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == ">" and secondSubSeq[1] == "<" :
					d = ( int(secondSubSeq[0]) ) - ( int(firstSubSeq[0]) + 6)
					if Interdistance_maxValue >= d >= 0 :
						ER.append(["ER "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == "<" and secondSubSeq[1] == ">" :
					d = ( int(secondSubSeq[0]) ) - ( int(firstSubSeq[0]) +  6 )
					if Interdistance_maxValue >= d >= 0 :
						IR.append(["IR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])

			
			if factorTranscription == "ARF2" :
				if firstSubSeq[1] == secondSubSeq[1] :
					d = ( int(secondSubSeq[0]) + 2 ) - ( int(firstSubSeq[0]) + lenMotif - 2)
					if Interdistance_maxValue >= d >= 0 :
						DR.append(["DR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == ">" and secondSubSeq[1] == "<" :
					d = ( int(secondSubSeq[0]) + 2 ) - ( int(firstSubSeq[0]) + lenMotif - 2 )
					if Interdistance_maxValue >= d >= 0 :
						ER.append(["ER "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == "<" and secondSubSeq[1] == ">" :
					d = ( int(secondSubSeq[0]) + 2 ) - ( int(firstSubSeq[0]) + lenMotif - 2 )
					if Interdistance_maxValue >= d >= 0 :
						IR.append(["IR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])

			if factorTranscription == "ARF5_OMalley_extended3Prime":
				if firstSubSeq[1] == ">" and secondSubSeq[1] == ">" :
					d = (int(secondSubSeq[0])  ) - (int(firstSubSeq[0]) + 6  )
					if Interdistance_maxValue >= d >= 0 :
						DR.append(["DR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == "<" and secondSubSeq[1] == "<" :
					d = (int(secondSubSeq[0])  ) - (int(firstSubSeq[0]) + 6 )
					if Interdistance_maxValue >= d >= 0 :
						DR.append(["DR " + str(d), firstSubSeq[1], firstSubSeq[2], firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == ">" and secondSubSeq[1] == "<" :
					d = (int(secondSubSeq[0])  ) - (int(firstSubSeq[0]) + 6 )
					if Interdistance_maxValue >= d >= 0 :
						ER.append(["ER "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == "<" and secondSubSeq[1] == ">" :
					d = (int(secondSubSeq[0])  ) - (int(firstSubSeq[0]) + 6  )
					if Interdistance_maxValue >= d >= 0 :
						IR.append(["IR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
			if factorTranscription == "ARF5" :
				if firstSubSeq[1] == ">" and secondSubSeq[1] == ">" :
					d = (int(secondSubSeq[0])  ) - (int(firstSubSeq[0]) + 6  )
					if Interdistance_maxValue >= d >= 0 :
						DR.append(["DR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == "<" and secondSubSeq[1] == "<" :
					d = (int(secondSubSeq[0])  ) - (int(firstSubSeq[0]) + 6 )
					if Interdistance_maxValue >= d >= 0 :
						DR.append(["DR " + str(d), firstSubSeq[1], firstSubSeq[2], firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == ">" and secondSubSeq[1] == "<" :
					d = (int(secondSubSeq[0])  ) - (int(firstSubSeq[0]) + 6 )
					if Interdistance_maxValue >= d >= 0 :
						ER.append(["ER "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
				if firstSubSeq[1] == "<" and secondSubSeq[1] == ">" :
					d = (int(secondSubSeq[0])  ) - (int(firstSubSeq[0]) + 6  )
					if Interdistance_maxValue >= d >= 0 :
						IR.append(["IR "+str(d),firstSubSeq[1],firstSubSeq[2],firstSubSeq[3],firstSubSeq[4],firstSubSeq[5],secondSubSeq[1],secondSubSeq[2],secondSubSeq[3],secondSubSeq[4],secondSubSeq[5],firstSubSeq[2]+secondSubSeq[2]])
		
	return(DR,ER,IR)	
	

# These 3 lines allows to retrieve the matrix from the file
F = open(MatrixFile,"r")
matrix = F.read().replace("\r","\n") + "\n"
F.close()

# These 3 lines allows to retrieve all the individual frequency values from the matrix and put them in order into a list
import re
num = re.compile(r"([+-]?\d+[.,]\d+)")
Mdata = num.findall(matrix)

matF, lenMotif = get_score_matrix(Mdata,matrixType,pseudoCount)

# The following line allows to produce the reversed matrix
'''if we take the example given before : A T G C
			Position 1:      0.4444  0.155  0.654   0.645
			Position 2:      0.1645  0.1565 0.21614 0.16456
Now, we can notice that scores change between the positions 1 and 2, and between A and T, and between G and C.
So we can calculate with this reverse matrix, the score of the complementary strand.
'''
matRev = list(reversed(matF))	
	
# This line allows to retrieve all the sequences from the fasta file
sequences = SeqIO.to_dict(SeqIO.parse(FastaFile, "fasta"))

print "  There are %s sequence(s) to analyze"%(len(sequences)) 

all_data = []
sequences_ = ""
explanation = ["Position of the T of the \"TGTCNN\" or of the N of the \"NNGACA\" if the motif is on the complementary strand" , "strand", "score", "binding site"]

# We look at all the fasta sequences:
for s in sorted(sequences):
	# We will store in this list all the best scores (see the threshold after) found for subsequences of one sequence
	#if type(threshold) is list:
	good_score_positions = [] 
	DR = []
	ER = []
	IR = []
	
	# This line allows to retrieve the DNA sequence
	seq = sequences[s].seq
	seq2 = ""
	seq_id = sequences[s].id
	a = 0
	p = document.add_paragraph("")
	p.add_run("Command line: "+str(sys.argv).replace("[","").replace("]","").replace("'","").replace(",",""))
	
	p = document.add_paragraph("")
	p.add_run("Legend: ").bold = True
	run = p.add_run("strand + ")
	run.font.color.rgb = RGBColor(0xFF, 0x00, 0xFF)
	run = p.add_run("strand - ")
	run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
	run = p.add_run("both overlap")
	run.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
	p = document.add_paragraph("")
	# We look at each sub-sequences of the whole sequence. Each sub-sequence has the same length that the matrix length.
	couleur1 = [0] * len(seq)
	couleur2 = [0] * len(seq)
	for c in range(len(seq) - (lenMotif -1)):
		strandPos = seq[c:c+lenMotif].upper()
		test = 0
		for nu in strandPos :
			if nu not in ["A","C","G","T"]:
				test = 1
		if test == 1:
			score = "NA"
		else :
			index = 0
			#These lines allows to calculate a score for one sub-sequence
			scoreStrandPos = 0
			scoreStrandNeg = 0 
			while index < lenMotif:
				if strandPos[index] == 'A':
					scoreStrandPos = scoreStrandPos + matF[index*4]
					scoreStrandNeg = scoreStrandNeg + matRev[index*4]
				elif strandPos[index] == 'C':
					scoreStrandPos = scoreStrandPos + matF[index*4+1]
					scoreStrandNeg = scoreStrandNeg + matRev[index*4+1]
				elif strandPos[index] == 'G':
					scoreStrandPos = scoreStrandPos + matF[index*4+2]
					scoreStrandNeg = scoreStrandNeg + matRev[index*4+2]
				elif strandPos[index] == 'T':
					scoreStrandPos = scoreStrandPos + matF[index*4+3]
					scoreStrandNeg = scoreStrandNeg + matRev[index*4+3]			
				index += 1

			#These lines allows to retrieve the position and the strand where there is a predicted binding site. 
			#You can change the threshold.
			#print("strandPos : ",strandPos)
			#print("scoreStrandPos : ",scoreStrandPos)
			if scoreStrandPos > threshold:
				#print("scoreStrandPos : ",scoreStrandPos)
				if factorTranscription == "ARF2_21bases" :
					couleur1[c+2] = 1
					couleur1[c+3] = 1
					couleur1[c+4] = 1
					couleur1[c+5] = 1
					couleur1[c+6] = 1
					couleur1[c+7] = 1
					rest = int((c+2)/65) + 1
					column = c+2 - int((c+2)/65) * 65
					good_score_positions.append([c+2,">",round(scoreStrandPos,2),str(strandPos[1:7]),rest,column])
				if factorTranscription == "ARF2_MEME1500" :
					couleur1[c+2] = 1
					couleur1[c+3] = 1
					couleur1[c+4] = 1
					couleur1[c+5] = 1
					couleur1[c+6] = 1
					couleur1[c+7] = 1
					rest = int((c+2)/65) + 1
					column = c+2 - int((c+2)/65) * 65
					good_score_positions.append([c+2,">",round(scoreStrandPos,2),str(strandPos[1:7]),rest,column])
				if factorTranscription == "ARF2" :
					couleur1[c+3] = 1
					couleur1[c+4] = 1
					couleur1[c+5] = 1
					couleur1[c+6] = 1
					couleur1[c+7] = 1
					couleur1[c+8] = 1
					rest = int((c+3)/65) + 1
					column = c+3 - int((c+3)/65) * 65
					good_score_positions.append([c+3,">",round(scoreStrandPos,2),str(strandPos[2:8]),rest,column])
				if factorTranscription == "ARF5" :
					couleur1[c+4] = 1
					couleur1[c+5] = 1
					couleur1[c+6] = 1
					couleur1[c+7] = 1
					couleur1[c+8] = 1
					couleur1[c+9] = 1
					rest = int((c+4)/65) + 1
					column = c+4 - int((c+4)/65) * 65
					good_score_positions.append([c+4,">",round(scoreStrandPos,2),str(strandPos[3:9]),rest,column])
				if factorTranscription == "ARF5_OMalley_extended3Prime":
					couleur1[c+4] = 1
					couleur1[c+5] = 1
					couleur1[c+6] = 1
					couleur1[c+7] = 1
					couleur1[c+8] = 1
					couleur1[c+9] = 1
					rest = int((c+4)/65) + 1
					column = c+4 - int((c+4)/65) * 65
					good_score_positions.append([c+4,">",round(scoreStrandPos,2),str(strandPos[3:9]),rest,column])
			if scoreStrandNeg > threshold:
				if factorTranscription == "ARF2_21bases" :
					couleur2[c+14] = 2
					couleur2[c+15] = 2
					couleur2[c+16] = 2
					couleur2[c+17] = 2
					couleur2[c+18] = 2
					couleur2[c+19] = 2
					rest = int((c+14)/65) + 1
					column = c+14 - int((c+14)/65) * 65
					good_score_positions.append([c+14,"<",round(scoreStrandNeg,2),str(strandPos[13:23]),rest,column])
				if factorTranscription == "ARF2_MEME1500" :
					couleur2[c+5] = 2
					couleur2[c+6] = 2
					couleur2[c+7] = 2
					couleur2[c+8] = 2
					couleur2[c+9] = 2
					couleur2[c+10] = 2
					rest = int((c+5)/65) + 1
					column = c+5 - int((c+5)/65) * 65
					good_score_positions.append([c+5,"<",round(scoreStrandNeg,2),str(strandPos[4:10]),rest,column])
				if factorTranscription == "ARF2" :
					couleur2[c+3] = 2
					couleur2[c+4] = 2
					couleur2[c+5] = 2
					couleur2[c+6] = 2
					couleur2[c+7] = 2
					couleur2[c+8] = 2
					rest = int((c+3)/65) + 1
					column = c+3 - int((c+3)/65) * 65
					good_score_positions.append([c+3,"<",round(scoreStrandNeg,2),str(strandPos[2:8]),rest,column])
				if factorTranscription == "ARF5_OMalley_extended3Prime" :
					couleur2[c+3] = 2
					couleur2[c+4] = 2
					couleur2[c+5] = 2
					couleur2[c+6] = 2
					couleur2[c+7] = 2
					couleur2[c+8] = 2
					rest = int((c+3)/65) + 1
					column = c+3 - int((c+3)/65) * 65
					good_score_positions.append([c+3,"<",round(scoreStrandNeg,2),str(strandPos[2:8]),rest,column])
				if factorTranscription == "ARF5" :
					couleur1[c+2] = 2
					couleur1[c+3] = 2
					couleur1[c+4] = 2
					couleur1[c+5] = 2
					couleur1[c+6] = 2
					couleur1[c+7] = 2
					rest = int((c+2)/65) + 1
					column = c+2 - int((c+2)/65) * 65
					good_score_positions.append([c+2,"<",round(scoreStrandNeg,2),str(strandPos[1:7]),rest,column])
	#print("good_score_positions : ",good_score_positions)				
	couleur = [x + y for x, y in zip(couleur1, couleur2)]
	#print("couleur : ",couleur)
	for c in range(len(seq) ):
		if couleur[c] == 0:
			run = p.add_run(seq[c-1:c])
			run.font.name = 'Courier 10 pitch'
			
		if couleur[c] == 1:
			run = p.add_run(seq[c-1:c])
			run.font.name = 'Courier 10 pitch'
			run.font.color.rgb = RGBColor(0xFF, 0x00, 0xFF)
			
		if couleur[c] == 2:
			run = p.add_run(seq[c-1:c])
			run.font.name = 'Courier 10 pitch'
			run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
			
		if couleur[c] == 3:
			run = p.add_run(seq[c-1:c])
			run.font.name = 'Courier 10 pitch'
			run.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)	
			
	run = document.add_paragraph().add_run("1            10           20            30            40            50            60")
	run.font.name = 'Courier 10 pitch'
	font = run.font
	font.size = Pt(8)
	
	seq2 = seq2 + seq[a:].upper()
	DR,ER,IR = interdistance_calcul(DR,ER,IR,good_score_positions)
	#print("DR : ",DR)
	#print("ER : ",ER)
	#print("IR : ",IR)
		
	#sequences_ = sequences_ + seq_id + "\n" + seq2 + "\n\n" + str(explanation) + "\nDR : " + str(DR) + "\nER : "+ str(ER) + "\nIR : " + str(IR) + "\n"  + "\n\nAll the sites:\n" + str(good_score_positions) + "\n\n\n"

########################################### About the main matrix #######################

''' The sens of the matrix is important: The positions are on the vertical sens and the bases are on the horizontal sens as described in the example.
separation between numbers can be spaces, tabulation, comas...

                                                                         Example :   A C G T
                                                                  position 1:           0.16456   0.21614       0.1565,0.1645
                                                                  position 2:           0.645; 0.654    0.155 |||||| 0.4444
                                                                                        ...
                                                                        '''
####################################################################################


###############################################################################################


	
########## get INTERDISTANCE VALUES for POSITIVE sets:

#sequences_ = get_interdist(matScore,matRev,FastaFile,threshold,factorTranscription,Interdistance_maxValue)

#output_presentation = "Here the binding sites are coloured. For example, with the sequence \n\"AGTAGTCAT TGT CAGATAGAAAGAAAGAG CAGACAAAAGGATCG\"\nBecause a binding site has a length of 6 bp, this means there is a first binding site: \nTGTCAGA, a second one: CAGATA and a last one: CAGACA.\nThe matrix used is: "+MatrixFile + "\n\n\n\n"
		
#text_file = open("localization_of_the_"+factorTranscription+"_binding_sites_on_"+FastaFile+".txt", "w")
#text_file.write(output_presentation+str(sequences_))
#text_file.close()


#p = document.add_paragraph('A plain paragraph having some ')
#p.add_run('bold').bold = True
#p.add_run(' and some ')
#p.add_run('italic.').italic = True

#document.add_heading('Heading, level 1', level=1)
#document.add_paragraph('Intense quote', style='IntenseQuote')


#document.add_paragraph(
    #'first item in ordered list', style='ListNumber'
#)
p = document.add_paragraph("")
p.add_run("DR table sorted by decreasing order for score sum: ").bold = True

table = document.add_table(rows=1, cols=6)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'DR'
hdr_cells[1].text = 'score sum'
hdr_cells[2].text = 'line'
hdr_cells[3].text = 'column'
hdr_cells[4].text = 'bs 1'
hdr_cells[5].text = 'bs 2'

DR = sorted(DR, key = operator.itemgetter(11), reverse=True)

for item in DR:
	row_cells = table.add_row().cells
	row_cells[0].text = str(item[0])
	row_cells[1].text = str(item[11])
	row_cells[2].text = str(item[4])
	row_cells[3].text = str(item[5])
	row_cells[4].text = str(item[3])
	row_cells[5].text = str(item[8])

document.add_page_break()

p = document.add_paragraph("")
p.add_run("ER table sorted by decreasing order for score sum: ").bold = True

table = document.add_table(rows=1, cols=6)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ER'
hdr_cells[1].text = 'score sum'
hdr_cells[2].text = 'line'
hdr_cells[3].text = 'column'
hdr_cells[4].text = 'bs 1'
hdr_cells[5].text = 'bs 2'

ER = sorted(ER, key = operator.itemgetter(11), reverse=True)

for item in ER:
	row_cells = table.add_row().cells
	row_cells[0].text = str(item[0])
	row_cells[1].text = str(item[11])
	row_cells[2].text = str(item[4])
	row_cells[3].text = str(item[5])
	row_cells[4].text = str(item[3])
	row_cells[5].text = str(item[8])

document.add_page_break()	

p = document.add_paragraph("")
p.add_run("IR table sorted by decreasing order for score sum: ").bold = True

table = document.add_table(rows=1, cols=6)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'IR'
hdr_cells[1].text = 'score sum'
hdr_cells[2].text = 'line'
hdr_cells[3].text = 'column'
hdr_cells[4].text = 'bs 1'
hdr_cells[5].text = 'bs 2'

IR = sorted(IR, key = operator.itemgetter(11), reverse=True)

for item in IR:
	row_cells = table.add_row().cells
	row_cells[0].text = str(item[0])
	row_cells[1].text = str(item[11])
	row_cells[2].text = str(item[4])
	row_cells[3].text = str(item[5])
	row_cells[4].text = str(item[3])
	row_cells[5].text = str(item[8])

document.add_page_break()	

p = document.add_paragraph("")
p.add_run("All bs table: ").bold = True

table = document.add_table(rows=1, cols=6)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Index'
hdr_cells[1].text = 'sens'
hdr_cells[2].text = 'score'
hdr_cells[3].text = 'sequence'
hdr_cells[4].text = 'line'
hdr_cells[5].text = 'column'

good_score_positions = sorted(good_score_positions, key=operator.itemgetter(2), reverse=True)

for item in good_score_positions:
	row_cells = table.add_row().cells
	row_cells[0].text = str(item[0])
	row_cells[1].text = str(item[1])
	row_cells[2].text = str(item[2])
	row_cells[3].text = str(item[3])
	row_cells[4].text = str(item[4])
	row_cells[5].text = str(item[5])

document.add_page_break()
FastaFile = FastaFile.replace("/","_")
MatrixFile = MatrixFile.replace("/","_")
document.save("localization_of_the_" + factorTranscription + "_binding_sites_on_" + FastaFile + ".docx")



